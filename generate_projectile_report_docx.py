#!/usr/bin/env python3
"""Generate 'Projectile_Motion_Report.docx' — a polished Word report on
2D projectile motion, with an embedded matplotlib trajectory plot."""

import math
import os

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Physics
# ---------------------------------------------------------------------------
V0, THETA_DEG, G = 25.0, 45.0, 9.81
theta = math.radians(THETA_DEG)

vx0 = V0 * math.cos(theta)
vy0 = V0 * math.sin(theta)
t_total = 2 * vy0 / G
h_max = vy0 ** 2 / (2 * G)
r_max = V0 ** 2 * math.sin(2 * theta) / G
t_apex = vy0 / G

N = 400
ts = [t_total * i / N for i in range(N + 1)]
xs = [vx0 * t for t in ts]
ys = [max(0.0, vy0 * t - 0.5 * G * t ** 2) for t in ts]

# ---------------------------------------------------------------------------
# Palette shared between the chart and the document
# ---------------------------------------------------------------------------
NAVY = "0B1F3A"
NAVY_RGB = RGBColor(0x0B, 0x1F, 0x3A)
ACCENT = "2EC4B6"
ACCENT_RGB = RGBColor(0x2E, 0xC4, 0xB6)
ACCENT_DARK_RGB = RGBColor(0x1E, 0x8F, 0x84)
GOLD = "F2A63D"
GOLD_RGB = RGBColor(0xF2, 0xA6, 0x3D)
RED = "E0473A"
RED_RGB = RGBColor(0xE0, 0x47, 0x3A)
TEXT_BODY_RGB = RGBColor(0x33, 0x3D, 0x4C)
MID_GRAY_RGB = RGBColor(0x6B, 0x74, 0x84)
LIGHT_GRAY = "E9EDF3"

PLOT_PATH = "trajectory_plot.png"

# ---------------------------------------------------------------------------
# 1. Trajectory plot
# ---------------------------------------------------------------------------
plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans"],
    "axes.edgecolor": "#333D4C",
    "axes.labelcolor": "#1B222D",
    "text.color": "#1B222D",
})

fig, ax = plt.subplots(figsize=(9, 5.2), dpi=300)

ax.plot(xs, ys, color="#0B1F3A", linewidth=2.6, zorder=3, label="Trajectory (v0 = 25 m/s, θ = 45°)")
ax.fill_between(xs, ys, 0, color="#2EC4B6", alpha=0.10, zorder=1)

# key points
ax.scatter([0], [0], s=90, color="#2EC4B6", edgecolor="white", linewidth=1.4, zorder=5, label="Launch Point (0, 0)")
ax.scatter([vx0 * t_apex], [h_max], s=110, color="#F2A63D", edgecolor="white", linewidth=1.4, zorder=5,
           label=f"Max Height ({vx0*t_apex:.2f}, {h_max:.2f})")
ax.scatter([r_max], [0], s=90, color="#E0473A", edgecolor="white", linewidth=1.4, zorder=5,
           label=f"Landing Point ({r_max:.2f}, 0)")

ax.annotate(f"H_max = {h_max:.2f} m", xy=(vx0 * t_apex, h_max), xytext=(vx0 * t_apex, h_max + 1.6),
            ha="center", fontsize=10, color="#1B222D", fontweight="bold")
ax.annotate(f"R_max = {r_max:.2f} m", xy=(r_max, 0), xytext=(r_max - 2, -2.6),
            ha="right", fontsize=10, color="#1B222D", fontweight="bold")

ax.axvline(vx0 * t_apex, color="#8A93A3", linestyle="--", linewidth=0.9, zorder=2)

ax.set_xlabel("Horizontal Distance, x (m)", fontsize=11.5, fontweight="bold")
ax.set_ylabel("Height, y (m)", fontsize=11.5, fontweight="bold")
ax.set_title("Projectile Trajectory — Height vs. Horizontal Distance", fontsize=13.5, fontweight="bold",
             color="#0B1F3A", pad=14)
ax.set_xlim(-1.5, r_max + 4)
ax.set_ylim(-3.2, h_max + 4)
ax.grid(True, linestyle="--", linewidth=0.6, alpha=0.55, zorder=0)
ax.set_axisbelow(True)
for spine in ("top", "right"):
    ax.spines[spine].set_visible(False)

legend = ax.legend(loc="upper right", frameon=True, fontsize=9.5, title="Key Points", title_fontsize=10)
legend.get_frame().set_edgecolor("#C7CFDB")
legend.get_frame().set_linewidth(0.8)
legend.get_title().set_fontweight("bold")

fig.tight_layout()
fig.savefig(PLOT_PATH, dpi=300, facecolor="white")
plt.close(fig)

# ---------------------------------------------------------------------------
# 2. Word document helpers
# ---------------------------------------------------------------------------

def set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_borders(cell, color="C7CFDB", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def style_cell_text(cell, text, bold=False, color=None, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1, color=NAVY_RGB, size=None, space_before=18, space_after=8):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    return h


def add_body_paragraph(doc, text, size=11, color=TEXT_BODY_RGB, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_formula_bullet(doc, label, formula, note=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{label}:  ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY_RGB
    r1.font.name = "Calibri"
    r2 = p.add_run(formula)
    r2.font.size = Pt(11.5)
    r2.font.name = "Cambria"
    r2.font.color.rgb = ACCENT_DARK_RGB
    r2.font.bold = True
    if note:
        r3 = p.add_run(f"   ({note})")
        r3.font.size = Pt(9.5)
        r3.font.italic = True
        r3.font.color.rgb = MID_GRAY_RGB
        r3.font.name = "Calibri"


def add_callout(doc, title, lines, accent_hex=ACCENT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "EAF3F7")
    set_cell_borders(cell, color=accent_hex, sz="10")
    cell.text = ""
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(4)
    rt = p_title.add_run(title)
    rt.font.bold = True
    rt.font.size = Pt(11)
    rt.font.color.rgb = ACCENT_DARK_RGB
    rt.font.name = "Calibri"
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"•  {line}")
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT_BODY_RGB
        r.font.name = "Calibri"
    # padding via cell margins
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", "150"), ("bottom", "150"), ("left", "200"), ("right", "200")):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tcPr.append(mar)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = MID_GRAY_RGB
    run.font.name = "Calibri"


def add_horizontal_rule(doc, color="2EC4B6"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# 3. Build the document
# ---------------------------------------------------------------------------
doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)

# ---- Title header ----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)
title_run = title_p.add_run("Physics Mechanics 101: 2D Projectile Motion Analysis")
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = NAVY_RGB
title_run.font.name = "Calibri"

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_p.paragraph_format.space_after = Pt(10)
subtitle_run = subtitle_p.add_run("Kinematics in Two Dimensions — Trajectory Modeling & Quantitative Analysis")
subtitle_run.font.size = Pt(13)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = ACCENT_DARK_RGB
subtitle_run.font.name = "Calibri"

add_horizontal_rule(doc, color=ACCENT)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_before = Pt(6)
meta_p.paragraph_format.space_after = Pt(24)
meta_run = meta_p.add_run("Prepared by: Physics Mechanics 101 Course Staff   |   Unit 2: Kinematics in Two Dimensions   |   Document Type: Technical Report")
meta_run.font.size = Pt(9.5)
meta_run.font.color.rgb = MID_GRAY_RGB
meta_run.font.name = "Calibri"

# ---- Section 1: Executive Summary & Overview ----
add_heading(doc, "1.  Executive Summary & Overview", level=1)

add_body_paragraph(
    doc,
    "One-dimensional kinematics describes motion confined to a single straight-line axis, where position, "
    "velocity, and acceleration are captured by one signed coordinate. Two-dimensional (2D) kinematics extends "
    "this framework to motion in a plane — most commonly a horizontal and a vertical axis — allowing objects "
    "to trace curved paths such as the parabolic arc of a thrown ball, launched projectile, or kicked object. "
    "Projectile motion is the canonical example: an object launched with some initial speed and angle above "
    "the horizontal, moving under the influence of gravity alone."
)

add_body_paragraph(
    doc,
    "The defining principle that makes 2D projectile motion tractable is the independence of the horizontal "
    "and vertical axes. Gravity acts exclusively along the vertical (y) direction, so the horizontal (x) "
    "velocity remains constant throughout the flight while the vertical velocity changes at a constant rate. "
    "This report derives the governing equations for both axes, computes the key trajectory metrics for a "
    "representative launch, and visualizes the resulting flight path."
)

# ---- Section 2: Governing Mathematical Model ----
add_heading(doc, "2.  Governing Mathematical Model", level=1)

add_body_paragraph(
    doc,
    "Under constant gravitational acceleration g and negligible air resistance, the horizontal and vertical "
    "motions decouple into independent 1D kinematics problems. The equations below describe position and "
    "velocity as functions of time t, given initial speed v0 and launch angle θ.",
    space_after=8,
)

add_formula_bullet(doc, "Horizontal Position", "x(t) = v0 · cos(θ) · t", "constant velocity, no acceleration")
add_formula_bullet(doc, "Vertical Position", "y(t) = v0 · sin(θ) · t − ½ g t²", "constant downward acceleration")
add_formula_bullet(doc, "Horizontal Velocity", "vx(t) = v0 · cos(θ)", "unchanging for the entire flight")
add_formula_bullet(doc, "Vertical Velocity", "vy(t) = v0 · sin(θ) − g t", "decreases, then reverses sign")
add_formula_bullet(doc, "Maximum Height", "H_max = (v0 · sin θ)² / (2g)", "occurs when vy = 0")
add_formula_bullet(doc, "Total Range", "R_max = v0² · sin(2θ) / g", "requires launch/landing at equal height")
add_formula_bullet(doc, "Time of Flight", "t_total = 2 · v0 · sin(θ) / g", "twice the time to reach apex")

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    doc,
    "Key Modeling Assumptions",
    [
        "Flat, level terrain — launch height and landing height are equal (y = 0 at both ends).",
        "Air resistance is negligible; drag and lift forces are not modeled.",
        "Gravitational acceleration g is constant in magnitude and direction throughout the flight.",
        "The projectile is treated as a point mass (no spin, no rotational dynamics).",
    ],
)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---- Section 3: Summary Table ----
add_heading(doc, "3.  Summary Table", level=1)
add_body_paragraph(
    doc,
    "Table 1 applies the governing equations above to a representative launch condition and reports the "
    "resulting trajectory metrics.",
    space_after=10,
)

headers = ["Parameter / Quantity", "Formula", "Calculated Value", "Units"]
rows = [
    ("Initial Velocity, v0", "— (given)", "25.00", "m/s"),
    ("Launch Angle, θ", "— (given)", "45.0", "degrees"),
    ("Horizontal Velocity, vx", "v0 · cos(θ)", "17.68", "m/s"),
    ("Vertical Velocity (initial), vy0", "v0 · sin(θ)", "17.68", "m/s"),
    ("Maximum Height, H_max", "(v0 sin θ)² / (2g)", "15.93", "m"),
    ("Total Range, R_max", "v0² sin(2θ) / g", "63.71", "m"),
    ("Time of Flight, t_total", "2 v0 sin(θ) / g", "3.60", "s"),
]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [Inches(1.9), Inches(1.9), Inches(1.3), Inches(0.9)]

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    style_cell_text(hdr_cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10.5,
                     align=WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_background(hdr_cells[i], NAVY)
    set_cell_borders(hdr_cells[i])
    hdr_cells[i].width = widths[i]

for r_idx, row in enumerate(rows):
    cells = table.add_row().cells
    band = "F4F7FA" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, val in enumerate(row):
        align = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
        bold = c_idx == 0
        style_cell_text(cells[c_idx], val, bold=bold, color=NAVY_RGB if c_idx == 0 else TEXT_BODY_RGB,
                         size=10, align=align)
        set_cell_background(cells[c_idx], band)
        set_cell_borders(cells[c_idx])
        cells[c_idx].width = widths[c_idx]

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ---- Section 4: Trajectory Plot & Visual Analysis ----
add_heading(doc, "4.  Trajectory Plot & Visual Analysis", level=1)

img_p = doc.add_paragraph()
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_run = img_p.add_run()
img_run.add_picture(PLOT_PATH, width=Inches(6))

add_caption(doc, "Figure 1. Projectile trajectory for v0 = 25 m/s, θ = 45°, g = 9.81 m/s², "
                 "with the launch point, apex (maximum height), and landing point marked.")

add_body_paragraph(
    doc,
    "The plotted trajectory is symmetric about the vertical line passing through the apex: the time to climb "
    "from launch to maximum height equals the time to descend from the apex back to ground level, and the "
    "impact speed at landing equals the launch speed (25 m/s), mirrored in direction below the horizontal. "
    "This symmetry is a direct consequence of constant gravitational acceleration acting on a launch and "
    "landing at equal elevation."
)

add_body_paragraph(
    doc,
    "The curvature of the path is parabolic rather than circular or linear because horizontal position grows "
    "linearly with time while vertical position is a quadratic function of time. Near the apex, the vertical "
    "velocity approaches zero and the path flattens, producing the characteristic rounded peak; away from the "
    "apex, the steeper vertical velocity component dominates and the path steepens, most visibly just before "
    "impact."
)

# ---- Section 5: Conclusion & Key Takeaways ----
add_heading(doc, "5.  Conclusion & Key Takeaways", level=1)

add_body_paragraph(
    doc,
    "This report has demonstrated that 2D projectile motion under constant gravity, with no air resistance, "
    "reduces cleanly to two independent 1D kinematics problems: uniform velocity along the horizontal axis "
    "and uniformly accelerated motion along the vertical axis. For the representative launch analyzed here "
    "(v0 = 25 m/s, θ = 45°), the projectile reaches a maximum height of approximately 15.93 m, travels a "
    "horizontal range of approximately 63.71 m, and remains airborne for approximately 3.60 s.",
    space_after=8,
)

takeaways = [
    "Horizontal and vertical motion are independent — gravity only affects the vertical component.",
    "A 45° launch angle maximizes range for a given launch speed on level ground.",
    "The trajectory is a symmetric parabola when launch and landing heights are equal.",
    "Real-world trajectories deviate from this ideal model once air resistance becomes significant.",
]
for item in takeaways:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(item)
    r.font.size = Pt(11)
    r.font.color.rgb = TEXT_BODY_RGB
    r.font.name = "Calibri"

doc.save("Projectile_Motion_Report.docx")

if os.path.exists(PLOT_PATH):
    os.remove(PLOT_PATH)

print("Saved Projectile_Motion_Report.docx")
